import bottle
from model import *
import docx

PISKOTEK_UPORABNISKO_IME = "uporabnisko_ime"
SKRIVNOST = "to je ena skrivnost"

@bottle.get("/")
def zacetna_stran():
    return bottle.template("prva_stran.html", napaka=None)


@bottle.post("/prijava/")
def prijava_post():
    uporabnisko_ime = bottle.request.forms.getunicode("uporabnisko_ime")
    geslo_v_cistopisu = bottle.request.forms.getunicode("geslo")
    if not uporabnisko_ime:
        return bottle.template("prva_stran.html", napaka="Vnesi uporabniško ime!")
    else:
        try:
            Uporabnik.prijava(uporabnisko_ime, geslo_v_cistopisu)
            bottle.response.set_cookie(PISKOTEK_UPORABNISKO_IME, uporabnisko_ime, path="/", secret=SKRIVNOST)
            bottle.redirect("/pozdrav_uporabnika/")
        except ValueError as e:
            return bottle.template("prva_stran.html", napaka=e)


@bottle.get("/pozdrav_uporabnika/")
def pozdrav_uporabnika():
    uporabnisko_ime = bottle.request.get_cookie(PISKOTEK_UPORABNISKO_IME, secret=SKRIVNOST)
    return bottle.template("pozdrav_uporabnika.html", uporabnisko_ime=uporabnisko_ime)


@bottle.get("/registracija/")
def registracija():
    return bottle.template("registracija.html", napake=[])


@bottle.post("/registracija/")
def registracija_post():
    uporabnisko_ime = bottle.request.forms.getunicode("uporabnisko_ime")
    geslo_v_cistopisu = bottle.request.forms.getunicode("geslo")
    
    if not uporabnisko_ime:
        return bottle.template("registracija.html", napake="Vnesi uporabniško ime!")
    elif not geslo_v_cistopisu:
        return bottle.template("registracija.html", napake="Vnesi geslo!")
    else:
        try:
            Uporabnik.registracija(uporabnisko_ime, geslo_v_cistopisu)
            bottle.response.set_cookie(PISKOTEK_UPORABNISKO_IME, uporabnisko_ime, path="/", secret=SKRIVNOST)
            bottle.redirect("/")
        except ValueError as e:
            return bottle.template("registracija.html", napake=e)


@bottle.get("/nov_test_osnova/")
def nov_test__osnova():
    return bottle.template("nov_test_osnova.html", napaka = None)


@bottle.post("/nov_test_osnova/")
def glava_testa():
    uporabnisko_ime = bottle.request.get_cookie(PISKOTEK_UPORABNISKO_IME, secret=SKRIVNOST)
    uporabnik = Uporabnik.iz_datoteke(uporabnisko_ime)

    predmet = bottle.request.forms.getunicode("predmet")
    letnik = bottle.request.forms.getunicode("letnik")

    if not predmet or not letnik:
        napaka = "Izpolniti morate vsa polja!"
        return bottle.template("nov_test_osnova.html", napaka = napaka)

    try:
        st_ucencev = int(bottle.request.forms.getunicode("st_ucencev"))
        st_nalog = int(bottle.request.forms.getunicode("st_nalog"))
    except ValueError:
        napaka = "Število učencev in število nalog morata biti številki!"
        return bottle.template("nov_test_osnova.html", napaka = napaka)
    
    if not st_nalog or not st_ucencev:
        napaka = "Število učencev in število nalog ne smeta biti enaki nič!"
        return bottle.template("nov_test_osnova.html", napaka = napaka)

    index_testa = uporabnik.nov_test(Test(uporabnik.uporabnisko_ime, predmet, letnik, st_ucencev, st_nalog))
    slovar_nalog = uporabnik.seznam_testov[index_testa].slovar_nalog

    uporabnik.v_datoteko()

    return bottle.template("nov_test_naloga.html", st_nalog=st_nalog, index_testa=index_testa, slovar_nalog=slovar_nalog, napaka=None)


@bottle.post("/uredi_besedilo/")
def uredi_besedilo():
    uporabnisko_ime = bottle.request.get_cookie(PISKOTEK_UPORABNISKO_IME, secret=SKRIVNOST)
    uporabnik = Uporabnik.iz_datoteke(uporabnisko_ime)

    izpolnjena_naloga = int(bottle.request.forms.getunicode("izpolnjena_naloga"))
    index_testa = int(bottle.request.forms.getunicode("index_testa"))
    st_nalog = int(bottle.request.forms.getunicode("st_nalog"))
    besedilo = bottle.request.forms.getunicode("besedilo")

    slovar_nalog = uporabnik.seznam_testov[index_testa].slovar_nalog
    naloga = slovar_nalog[izpolnjena_naloga]

    if besedilo.count("#") == 0:
        napaka_besedilo = "V besedilo niste vnesli nobenega spremenljivega podatka!"
        return bottle.template("nov_test_naloga.html", st_nalog=st_nalog, index_testa=index_testa, slovar_nalog=slovar_nalog, napaka=napaka_besedilo)

    if not naloga.preveri_besedilo(besedilo):
        napaka_besedilo = "Besedilo ste napisali narobe! Podatke oštevilčite z naravnimi števili po vrsti!"
        return bottle.template("nov_test_naloga.html", st_nalog=st_nalog, index_testa=index_testa, slovar_nalog=slovar_nalog, napaka=napaka_besedilo)
    
    naloga.spremeni_besedilo(besedilo)
    uporabnik.v_datoteko()

    return bottle.template("nov_test_naloga.html", st_nalog=st_nalog, index_testa=index_testa, slovar_nalog=slovar_nalog, napaka=None)


@bottle.post("/uredi_podatke/")
def uredi_podatke():
    uporabnisko_ime = bottle.request.get_cookie(PISKOTEK_UPORABNISKO_IME, secret=SKRIVNOST)
    uporabnik = Uporabnik.iz_datoteke(uporabnisko_ime)

    izpolnjena_naloga = int(bottle.request.forms.getunicode("izpolnjena_naloga"))
    index_testa = int(bottle.request.forms.getunicode("index_testa"))
    st_nalog = int(bottle.request.forms.getunicode("st_nalog"))

    slovar_nalog = uporabnik.seznam_testov[index_testa].slovar_nalog
    naloga = uporabnik.seznam_testov[index_testa].slovar_nalog[izpolnjena_naloga]
    st_podatkov = naloga.st_podatkov

    podatki = {}
    for i in range(st_podatkov):
        b = bottle.request.forms.getunicode(f"answer{i+1}")
        z = bottle.request.forms.getunicode(f"answer{i+1}_zacetek")
        k = bottle.request.forms.getunicode(f"answer{i+1}_konec")
        podatki[f"#{i+1}"] = (b, z, k)

        if int(k) <= int(z):
            napaka_resitev = "Narobe ste izbrali interval!"
            return bottle.template("nov_test_naloga.html", st_nalog=st_nalog, index_testa=index_testa, slovar_nalog=slovar_nalog, napaka=napaka_resitev)
        elif b == "N" and int(z) < 0:
            napaka_resitev = "Narobe ste izbrali interval!"
            return bottle.template("nov_test_naloga.html", st_nalog=st_nalog, index_testa=index_testa, slovar_nalog=slovar_nalog, napaka=napaka_resitev)
    naloga.spremeni_slovar_baz(podatki)

    formula_resitve = bottle.request.forms.getunicode("formula_resitve")
    if not formula_resitve or formula_resitve.count('#') == 0:
        napaka_resitev = "Niste vnesli formule za rešitev!"
        return bottle.template("nov_test_naloga.html", st_nalog=st_nalog, index_testa=index_testa, slovar_nalog=slovar_nalog, napaka=napaka_resitev)
    else:
        naloga.spremeni_formulo(formula_resitve)
        naloga.ustvari_razlicice()
        # try:
        #     naloga.spremeni_formulo(formula_resitve)
        #     naloga.ustvari_razlicice()
        # except SyntaxError:
        #     napaka_resitve = "Formulo ste napisali narobe!"
        #     return bottle.template("nov_test_naloga.html", st_nalog=st_nalog, index_testa=index_testa, slovar_nalog=slovar_nalog, napaka=napaka_resitve)

    naloga.spremeni_stanje('KN')
    uporabnik.seznam_testov[index_testa].posodobi_stanje()
    uporabnik.v_datoteko()
    return bottle.template("nov_test_naloga.html", st_nalog=st_nalog, index_testa=index_testa, slovar_nalog=slovar_nalog, napaka = None)


@bottle.post("/test/")
def test():
    uporabnisko_ime = bottle.request.get_cookie(PISKOTEK_UPORABNISKO_IME, secret=SKRIVNOST)
    uporabnik = Uporabnik.iz_datoteke(uporabnisko_ime)
    index_testa = int(bottle.request.forms.getunicode("index_testa"))

    test = uporabnik.seznam_testov[index_testa]
    test.posodobi_stanje()
    slovar_nalog = test.slovar_nalog
    uporabnik.v_datoteko()

    dokument = docx.Document()
    dokument.add_heading(f"Rešitve {test.predmet}, {test.letnik}", level = 1)
    for i in slovar_nalog:
        naloga = slovar_nalog[i]
        dokument.add_heading(f"{i + 1}. naloga", level = 2)
        for j in range(test.st_razlicic):
            razlicica = naloga.seznam_razlicic[j]
            resitev = razlicica.resitev                
            dokument.add_paragraph(f"učenec {j + 1}: {resitev}")
        dokument.add_paragraph()
    dokument.save(f'testi/{uporabnisko_ime}_{test.predmet}_{test.letnik}_resitve.docx')
        
    for i in range(test.st_razlicic):
        dokument = docx.Document()
        dokument.add_heading(test.glava, level = 1)
        dokument.add_heading(f"učenec {i + 1}")
        for j in slovar_nalog:
            seznam_razlicic = slovar_nalog[j].seznam_razlicic
            razlicica = seznam_razlicic[i]
            besedilo = razlicica.besedilo

            dokument.add_heading(f"{j + 1}. naloga", level = 3)
            dokument.add_paragraph(f"{besedilo}")
            dokument.add_paragraph()
            dokument.add_paragraph()
        dokument.save(f'testi/{uporabnisko_ime}_{test.predmet}_{test.letnik}_ucenec{i + 1}.docx')

    return bottle.template("test.html" , test=test, uporabnisko_ime=uporabnisko_ime)


@bottle.get("/arhiv/")
def arhiv_testov():
    uporabnisko_ime = bottle.request.get_cookie(PISKOTEK_UPORABNISKO_IME, secret=SKRIVNOST)
    uporabnik = Uporabnik.iz_datoteke(uporabnisko_ime)
    seznam_testov = uporabnik.seznam_testov
    return bottle.template("arhiv.html", seznam_testov = seznam_testov)


@bottle.post("/nedokoncan_test/")
def nedokoncan_test():
    uporabnisko_ime = bottle.request.get_cookie(PISKOTEK_UPORABNISKO_IME, secret=SKRIVNOST)
    uporabnik = Uporabnik.iz_datoteke(uporabnisko_ime) 

    # izpolnjena_naloga = int(bottle.request.forms.getunicode("izpolnjena_naloga"))
    index_testa = int(bottle.request.forms.getunicode("index_testa"))
    st_nalog = int(bottle.request.forms.getunicode("st_nalog"))
    slovar_nalog = uporabnik.seznam_testov[index_testa].slovar_nalog

    return bottle.template("nov_test_naloga.html", st_nalog=st_nalog, index_testa=index_testa, slovar_nalog=slovar_nalog, napaka=None)


@bottle.post("/koncan_test/")
def koncan_test():
    uporabnisko_ime = bottle.request.get_cookie(PISKOTEK_UPORABNISKO_IME, secret=SKRIVNOST)
    uporabnik = Uporabnik.iz_datoteke(uporabnisko_ime)
    index_testa = int(bottle.request.forms.getunicode("index_testa"))

    test = uporabnik.seznam_testov[index_testa]

    return bottle.template("test.html", test=test, uporabnisko_ime=uporabnisko_ime)


@bottle.post("/odjava/")
def odjava():
    bottle.response.delete_cookie(PISKOTEK_UPORABNISKO_IME)
    bottle.redirect("/")
    

@bottle.route('/static/<filename:path>')
def send_static(filename):
    return bottle.static_file(filename, root='static')


@bottle.route('/testi/<filename:path>')
def send_static(filename):
    return bottle.static_file(filename, root='testi')


if __name__ == "__main__":
    bottle.run(debug=True, reloader=True)